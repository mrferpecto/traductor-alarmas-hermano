import streamlit as st
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from deep_translator import GoogleTranslator
import os
import subprocess

# --- CONFIGURACIÓN ---
st.set_page_config(page_title="Traductor Alarmas V5", page_icon="🔧")

st.title("🔧 Traductor V5 (Ajuste Inteligente)")
st.markdown("""
**Correcciones aplicadas:**
1. **Eliminación de huecos:** Se resetea el espaciado entre párrafos.
2. **Auto-Escalado:** Si el texto traducido es muy largo, se reduce la letra para que quepa.
3. **Alineación Izquierda:** Para evitar palabras cortadas y huecos raros.
""")

# --- FUNCIONES ---

def traducir_texto(texto, idioma_destino):
    if not texto or len(texto) < 2:
        return texto
    # Limpiamos saltos de línea extraños dentro de frases
    texto_limpio = texto.replace('-\n', '').replace('\n', ' ')
    try:
        return GoogleTranslator(source='auto', target=idioma_destino).translate(texto_limpio)
    except:
        return texto

def aplicar_estilo_seguro(run_origen, run_destino, factor_largo=1.0):
    """
    Copia el estilo pero reduce la fuente si el texto nuevo es mucho más largo
    para evitar que se corte o rompa la maquetación.
    """
    try:
        # Copiar atributos básicos
        run_destino.bold = run_origen.bold
        run_destino.italic = run_origen.italic
        run_destino.font.name = run_origen.font.name
        
        # Copiar color
        if run_origen.font.color and run_origen.font.color.rgb:
            run_destino.font.color.rgb = run_origen.font.color.rgb
        
        # Lógica de Auto-Ajuste de tamaño
        tamano_base = run_origen.font.size
        if tamano_base:
            # Si el texto creció más de un 20%, reducimos la fuente un poco
            if factor_largo > 1.2:
                nuevo_tamano = tamano_base.pt * 0.85 # Reducir 15%
                run_destino.font.size = Pt(max(7, nuevo_tamano)) # Nunca menos de 7pt
            else:
                run_destino.font.size = tamano_base
    except:
        pass

def limpiar_parrafo(parrafo):
    """Elimina espaciados gigantescos (brutales) entre párrafos"""
    pf = parrafo.paragraph_format
    pf.space_before = Pt(2) # Solo 2 puntos de espacio antes
    pf.space_after = Pt(2)  # Solo 2 puntos de espacio después
    pf.line_spacing = 1.1   # Interlineado sencillo pero legible
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT # Alineación izquierda segura

def procesar_bloque(bloque, lang_code):
    text_raw = bloque.text.strip()
    if not text_raw:
        return

    # 1. Analizar estilo original (usamos el primer run con contenido)
    estilo_ref = None
    len_original = len(text_raw)
    
    for run in bloque.runs:
        if run.text.strip():
            estilo_ref = run
            break
            
    if not estilo_ref:
        return

    # 2. Traducir
    traduccion = traducir_texto(text_raw, lang_code)
    len_traduccion = len(traduccion)
    
    # Calculamos cuánto ha crecido el texto
    factor_crecimiento = len_traduccion / len_original if len_original > 0 else 1

    # 3. Reemplazar contenido
    bloque.clear()
    nuevo_run = bloque.add_run(traduccion)
    
    # 4. Aplicar estilo con inteligencia
    aplicar_estilo_seguro(estilo_ref, nuevo_run, factor_crecimiento)
    
    # 5. Arreglar espaciados del párrafo contenedor
    if hasattr(bloque, 'paragraph_format'):
        limpiar_parrafo(bloque)
    else:
        # Si es una celda, a veces hay que acceder a sus párrafos internos
        pass 

def procesar_documento(input_pdf, lang_code):
    base_name = os.path.splitext(input_pdf)[0]
    temp_docx = f"{base_name}_temp.docx"
    
    # 1. PDF -> DOCX
    cv = Converter(input_pdf)
    # intersection_X_tolerance alto ayuda a fusionar columnas rotas
    cv.convert(temp_docx, start=0, end=None)
    cv.close()

    # 2. Procesar DOCX
    doc = Document(temp_docx)
    total_bloques = len(doc.paragraphs) + len(doc.tables)
    bar = st.progress(0)
    contador = 0
    
    # Párrafos sueltos
    for para in doc.paragraphs:
        procesar_bloque(para, lang_code)
        contador += 1
        if contador % 10 == 0: bar.progress(min(contador/total_bloques, 0.8))

    # Tablas (Iteramos celda por celda)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    procesar_bloque(p, lang_code)
                    # Forzar limpieza de espacios en celdas de tabla también
                    limpiar_parrafo(p)
        contador += 1

    bar.progress(0.9)
    doc.save(temp_docx)
    
    # 3. DOCX -> PDF
    st.info("Generando PDF final optimizado...")
    cwd = os.getcwd()
    cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', temp_docx, '--outdir', cwd]
    subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    
    pdf_final = f"{base_name}_temp.pdf"
    bar.progress(1.0)
    
    return pdf_final, temp_docx

# --- INTERFAZ ---

uploaded_file = st.file_uploader("Sube PDF (Versión V5 Reparada)", type=["pdf"])
idiomas = {"Alemán": "de", "Inglés": "en", "Francés": "fr", "Holandés": "nl", "Italiano": "it"}
target = st.selectbox("Idioma:", list(idiomas.keys()))

if uploaded_file and st.button("TRADUCIR Y ARREGLAR", type="primary"):
    with st.spinner('Procesando... Aplicando reducción de espacios y ajuste de texto...'):
        try:
            nombre_entrada = "entrada.pdf"
            with open(nombre_entrada, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            pdf_resultante, docx_temp = procesar_documento(nombre_entrada, idiomas[target])
            
            if os.path.exists(pdf_resultante):
                st.success("¡Listo!")
                with open(pdf_resultante, "rb") as f:
                    st.download_button("📥 DESCARGAR PDF", f, file_name=f"Contrato_{target}.pdf")
                # Limpieza
                try:
                    os.remove(nombre_entrada)
                    os.remove(docx_temp)
                    os.remove(pdf_resultante)
                except: pass
            else:
                st.error("Error al generar PDF.")
        except Exception as e:
            st.error(f"Error: {e}")
